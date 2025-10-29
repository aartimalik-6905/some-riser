# === 1. IMPORTS ===
# Standard libraries
import io
import os
import tempfile
import json
import time # Needed for retry

# Third-party libraries
import spacy
import pandas as pd
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import docx
from img2table.ocr import TesseractOCR
from img2table.document import Image as Img2TableImage
from spacy.lang.en.stop_words import STOP_WORDS
from flask import Flask, request, jsonify
from flask_cors import CORS
import google.generativeai as genai # AI brain

# --- 2. SETUP & FIXES ---

# FIX 1: Numba/img2table cache error [Errno 2]
numba_cache_dir = os.path.join(tempfile.gettempdir(), 'numba_cache')
if not os.path.exists(numba_cache_dir):
    os.makedirs(numba_cache_dir)
os.environ['NUMBA_CACHE_DIR'] = numba_cache_dir

# FIX 2: Tesseract PATH (Not strictly needed in Docker if installed via apt-get)
# The Dockerfile handles adding Tesseract to the system PATH

# --- 3. CONFIGURE THE AI BRAIN ---
# Load API key securely from environment
try:
    api_key = os.environ.get('API_KEY')
    if not api_key:
        print("CRITICAL ERROR: 'API_KEY' environment variable not set.")
        print("Please set it in your deployment platform's secrets/environment settings.")
        exit() # Stop the app if the key is missing
    genai.configure(api_key=api_key)
except Exception as e:
    print(f"Error configuring Google AI. {e}")
    exit()

# --- 4. LOAD SPACY MODEL ---
try:
    # Uses the model version specified in requirements.txt
    nlp = spacy.load("en_core_web_sm", disable=["lemmatizer"])
except OSError:
    print("Error: 'en_core_web_sm' model not found. Ensure it's installed via requirements.txt.")
    exit()

# --- 5. CREATE FLASK APP ---
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}}) # Allow all origins

# --- 6. HEALTH CHECK ROUTE --- ### ADDED BACK ###
@app.route("/")
def health_check():
    """Basic health check route for the deployment platform."""
    return jsonify({"status": "ok"})

# --- 7. SUMMARIZER FUNCTIONS ---

def summarize_document_gemini(text_content):
    """Analyzes document text using the Google AI (Gemini) API with retry."""
    if not text_content or not text_content.strip():
        return {"summary": "No text provided.", "keywords": [], "image_count": 0}

    # Use 'gemini-pro' as requested, though 'flash' is faster/cheaper
    model = genai.GenerativeModel('gemini-pro')
    prompt = f"""
    Analyze the following text and provide two things in a clean JSON format:
    1. "summary": A concise, abstractive summary (2-4 sentences).
    2. "keywords": A list of the 8-10 most important keywords/phrases.

    Output only the raw JSON object.

    Text:
    ---
    {text_content}
    ---
    """

    # --- RETRY LOGIC ADDED BACK ---
    for attempt in range(2): # Try up to 2 times (initial + 1 retry)
        try:
            print(f"Attempt {attempt + 1}: Calling Google AI...")
            response = model.generate_content(prompt)
            print("Google AI call successful.")

            print("Attempting to parse JSON...")
            json_text = response.text.strip().lstrip("```json").rstrip("```")
            ai_results = json.loads(json_text)
            print("JSON parsing successful.")

            return { # Return results on success
                "summary": ai_results.get("summary", "Summary failed."),
                "keywords": ai_results.get("keywords", []),
                "image_count": 0
            }

        except json.JSONDecodeError as json_e:
            print(f"Attempt {attempt + 1}: Error parsing JSON: {json_e}")
            raw_response = "N/A"
            if 'response' in locals() and hasattr(response, 'text'):
                raw_response = response.text
            print(f"Raw AI response: {raw_response}")
            # Don't retry JSON errors, return immediately
            return {"summary": "Error: Bad AI response.", "keywords": [], "image_count": 0}

        except Exception as e:
            print(f"Attempt {attempt + 1}: Error during Google AI call: {e}")
            print(f"Exception Type: {type(e).__name__}")
            raw_response = "N/A"
            if 'response' in locals() and hasattr(response, 'text'):
                raw_response = response.text
            print(f"Raw AI response: {raw_response}")

            # If this was the last attempt, return the error
            if attempt == 1:
                print("Max retries reached. Returning error.")
                return {"summary": "Error: AI call failed after retries.", "keywords": [], "image_count": 0}

            # Wait 2 seconds before retrying
            print("Waiting 2 seconds before retry...")
            time.sleep(2)
    # --- END RETRY LOGIC ---


def summarize_table_from_text(csv_text):
    """Analyzes CSV text content."""
    try:
        csv_file = io.StringIO(csv_text)
        df = pd.read_csv(csv_file)
        if df.empty:
            return {"error": "CSV file is empty or could not be read."}

        columns = df.columns.tolist()
        num_rows = df.shape[0]
        num_cols = df.shape[1]

        column_types_list = []
        for col, dtype in df.dtypes.items():
            col_type = 'Number' if 'int' in str(dtype) or 'float' in str(dtype) else 'Text'
            column_types_list.append(f"{col} ({col_type})")

        row_names = []
        row_header_name = ""
        # Check added: Ensure there's at least one column before accessing iloc[:, 0]
        if num_cols > 0 and df.iloc[:, 0].dtype == 'object':
            row_header_name = columns[0]
            row_names = df.iloc[:, 0].unique().tolist()[:5]

        sample_row = df.head(1).to_dict(orient='records')[0] if num_rows > 0 else {}

        insights = [
            f"Table has {num_rows} rows and {num_cols} columns.",
            f"Column types: {', '.join(column_types_list)}."
        ]
        if row_names:
            insights.append(f"First column '{row_header_name}' might be row headers.")
            insights.append(f"Examples: {', '.join(row_names)}.")
        # Check added: Ensure columns list is not empty before iterating
        if columns and any(kw in col.lower() for kw in ['sales', 'revenue', 'q1', 'q2'] for col in columns):
            insights.append("Table might track financial/quarterly data.")

        return {
            "columns": columns, "num_rows": num_rows, "num_cols": num_cols,
            "sample_row": sample_row, "column_types": column_types_list,
            "row_names": row_names, "insights": insights
        }
    except Exception as e:
        print(f"Error parsing table: {e}")
        return {"error": f"Could not parse table. Is it valid CSV? ({e})"}

# --- 8. API ROUTES ---

@app.route('/summarize-doc', methods=['POST'])
def handle_doc_summary():
    """Bridge for pasted text or .txt files."""
    try:
        data = request.json
        if 'text' not in data or not data['text'].strip():
            return jsonify({"error": "No text provided"}), 400
        text = data['text']
        summary_data = summarize_document_gemini(text)
        # Add placeholder for entities if your frontend expects it
        summary_data["entities"] = []
        return jsonify(summary_data)
    except Exception as e:
        print(f"Error in /summarize-doc: {e}")
        return jsonify({"error": "Server error analyzing text."}), 500

@app.route('/summarize-table', methods=['POST'])
def handle_table_summary():
    """Bridge for .csv files."""
    try:
        data = request.json
        if 'text' not in data:
            return jsonify({"error": "No CSV text provided"}), 400
        csv_text = data['text']
        summary_data = summarize_table_from_text(csv_text)
        return jsonify(summary_data)
    except Exception as e:
        print(f"Error in /summarize-table: {e}")
        return jsonify({"error": "Server error analyzing CSV."}), 500

@app.route('/summarize-image', methods=['POST'])
def handle_image_summary():
    """Bridge for .png/.jpg images. Uses img2table."""
    try:
        if 'image' not in request.files:
            return jsonify({"error": "No image file provided"}), 400
        file = request.files['image']
        img_bytes = file.read()

        # 1. Use img2table to find tables
        ocr = TesseractOCR(n_threads=1, lang="eng")
        doc = Img2TableImage(src=img_bytes)
        extracted_tables = doc.extract_tables(ocr=ocr, implicit_rows=True, borderless_tables=True, min_confidence=50)

        table_summary = {"error": "No table detected."}
        if extracted_tables:
            first_table_df = extracted_tables[0].dataframe
            # Check if dataframe is empty before proceeding
            if not first_table_df.empty:
                csv_text = first_table_df.to_csv(index=False)
                table_summary = summarize_table_from_text(csv_text)
            else:
                table_summary = {"error": "Detected table structure, but no data found."}


        # 2. Get Text Summary (pytesseract + Gemini)
        img_for_pytesseract = Image.open(io.BytesIO(img_bytes))
        extracted_text = pytesseract.image_to_string(img_for_pytesseract)

        if not extracted_text.strip() and not extracted_tables:
             return jsonify({"error": "Could not read text or detect table."}), 400

        doc_summary = summarize_document_gemini(extracted_text)
        if extracted_tables:
            doc_summary["summary"] = "Image identified as table. See insights below."
        doc_summary["entities"] = [] # Keep structure consistent

        # 3. Return Combined
        return jsonify({"doc_summary": doc_summary, "table_summary": table_summary})
    except Exception as e:
        print(f"Error processing image: {e}")
        return jsonify({"error": f"Could not analyze image. {e}"}), 500

@app.route('/summarize-mixed-doc', methods=['POST'])
def handle_mixed_doc():
    """Bridge for .pdf and .docx files."""
    try:
        if 'doc' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        file = request.files['doc']
        all_text = ""
        all_tables = []
        image_count = 0

        # 1. Read File
        if file.filename.endswith('.pdf'):
            # Ensure the stream is seekable for fitz
            file.stream.seek(0)
            with fitz.open(stream=file.stream.read(), filetype='pdf') as pdf_doc:
                for page in pdf_doc:
                    all_text += page.get_text() + "\n"
                    image_count += len(list(page.get_images(full=True)))
                    try:
                        # find_tables can sometimes error on complex pages
                        for table in page.find_tables():
                             # Check if extract method returns valid data
                             extracted_data = table.extract()
                             if extracted_data: # Only add non-empty tables
                                 all_tables.append(extracted_data)
                    except Exception as table_err:
                        print(f"Warning: Error finding tables on PDF page: {table_err}")

        elif file.filename.endswith('.docx'):
            # Ensure the stream is seekable for docx
            file.stream.seek(0)
            doc_obj = docx.Document(io.BytesIO(file.stream.read()))
            for para in doc_obj.paragraphs:
                all_text += para.text + "\n"
            for shape in doc_obj.inline_shapes:
                # Use attribute access, not index
                if hasattr(shape, 'type') and shape.type == 3: # 3 is wdInlineShapePicture
                    image_count += 1
            for table in doc_obj.tables:
                table_data = [[cell.text for cell in row.cells] for row in table.rows]
                if table_data: # Only add non-empty tables
                    all_tables.append(table_data)
        else:
             return jsonify({"error": "Unsupported file type"}), 400

        # 2. Summarize Text
        doc_summary_results = summarize_document_gemini(all_text)
        doc_summary_results["image_count"] = image_count
        doc_summary_results["entities"] = [] # Keep structure consistent

        # 3. Summarize First Table
        table_summary_results = {"error": "No tables found."}
        if all_tables:
            first_table_list = all_tables[0]
            # Check for valid table structure (at least header + 1 data row)
            if first_table_list and len(first_table_list) > 1:
                header = first_table_list[0]
                data = first_table_list[1:]
                # Check header has unique, non-empty values (basic validation)
                if header and data and len(header) > 0 and len(header) == len(set(filter(None, header))):
                    try:
                        df = pd.DataFrame(data, columns=header)
                        csv_text = df.to_csv(index=False)
                        table_summary_results = summarize_table_from_text(csv_text)
                    except Exception as df_err:
                         print(f"Error creating/summarizing DataFrame from table: {df_err}")
                         table_summary_results = {"error": "Found table, but failed to process data."}
                else:
                    table_summary_results = {"error": "Found table structure, but header/data invalid."}
            else:
                table_summary_results = {"error": "Found table, but it was empty/invalid."}

        # 4. Return Combined
        return jsonify({
            "doc_summary": doc_summary_results,
            "table_summary": table_summary_results
        })
    except Exception as e:
        print(f"Error processing mixed document: {e}")
        # Provide a more specific error if possible
        error_msg = f"Could not analyze file. {type(e).__name__}: {e}"
        return jsonify({"error": error_msg}), 500

# --- 9. RUN SERVER (Only for Local Testing) ---
if __name__ == '__main__':
    print("Starting Flask server for local testing...")
    port = int(os.environ.get("PORT", 5000))
    # Use debug=False for production readiness, listen on all interfaces
    app.run(debug=False, host='0.0.0.0', port=port)
